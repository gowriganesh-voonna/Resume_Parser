import logging
import os
import re
import shutil
import tempfile
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from PIL import Image

# Work around Paddle CPU runtime issues on Windows (oneDNN + PIR path).
# These must be set before importing paddle/paddleocr.
os.environ.setdefault("FLAGS_use_mkldnn", "0")
os.environ.setdefault("FLAGS_enable_pir_in_executor", "0")
os.environ.setdefault("FLAGS_enable_pir_api", "0")
os.environ.setdefault("FLAGS_use_pir_api", "0")

try:
    from paddleocr import PaddleOCR

    PADDLE_AVAILABLE = True
except ImportError:
    PADDLE_AVAILABLE = False
    print("PaddleOCR not available. Install with: pip install paddleocr paddlepaddle")

try:
    import pytesseract

    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

logger = logging.getLogger(__name__)


class FileProcessor:
    """Phase 3: File Processing Engine - Extract raw text from any resume format"""

    def __init__(self):
        self._configure_tesseract()
        self.ocr = None
        if PADDLE_AVAILABLE:
            self.ocr = self._init_paddle_ocr()

    def _configure_tesseract(self):
        """Resolve tesseract binary location on Windows when PATH is missing."""
        if not TESSERACT_AVAILABLE:
            return

        env_cmd = os.getenv("TESSERACT_CMD")
        local_app_data = os.getenv("LOCALAPPDATA")
        candidates = [
            env_cmd,
            shutil.which("tesseract"),
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            str(Path(local_app_data) / "Programs" / "Tesseract-OCR" / "tesseract.exe")
            if local_app_data
            else None,
        ]

        for candidate in candidates:
            if candidate and os.path.exists(candidate):
                pytesseract.pytesseract.tesseract_cmd = candidate
                logger.info(f"Using Tesseract binary: {candidate}")
                return

        logger.warning(
            "Tesseract binary not found. Set PATH or TESSERACT_CMD to tesseract.exe location."
        )

    def _init_paddle_ocr(self):
        """Initialize PaddleOCR across version-specific constructor signatures."""
        constructors = [
            {"use_angle_cls": True, "lang": "en", "show_log": False, "enable_mkldnn": False},
            {"use_angle_cls": True, "lang": "en", "show_log": False},
            {"use_angle_cls": True, "lang": "en", "enable_mkldnn": False},
            {"use_angle_cls": True, "lang": "en"},
        ]

        last_error = None
        for kwargs in constructors:
            try:
                ocr = PaddleOCR(**kwargs)
                logger.info(f"PaddleOCR initialized successfully with args: {kwargs}")
                return ocr
            except Exception as e:
                last_error = e

        logger.error(f"Failed to initialize PaddleOCR: {last_error}")
        return None

    def extract_text(self, file_path: str, file_type: Optional[str] = None) -> str:
        """
        Extract raw text from resume file based on file type
        Returns cleaned raw text
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Auto-detect file type if not provided
        if not file_type:
            file_type = self._detect_file_type(file_path)

        logger.info(f"Extracting text from {file_path} (type: {file_type})")

        # Extract based on file type
        if file_type in ["pdf", "application/pdf"]:
            text = self._extract_pdf_with_layout_awareness(file_path)
        elif file_type in ["docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            text = self._extract_from_docx_with_images(file_path)
        elif file_type in ["doc", "application/msword"]:
            text = self._extract_from_doc(file_path)
        elif file_type in ["txt", "text/plain"]:
            text = self._extract_from_txt(file_path)
        elif file_type in ["png", "jpg", "jpeg", "bmp", "tiff", "image/png", "image/jpeg"]:
            text = self._extract_from_image(file_path)
        else:
            # Try OCR as fallback for unknown types
            logger.warning(f"Unknown file type {file_type}, attempting OCR")
            text = self._extract_from_image(file_path)

        # Clean the extracted text
        cleaned_text = self._clean_text(text)

        return cleaned_text

    def _extract_pdf_with_layout_awareness(self, file_path: str) -> str:
        """
        PDF extraction with layout-aware strategy and robust fallback.
        This prevents pipeline failures from new extraction logic changes.
        """
        text_parts: list[str] = []
        try:
            from app.services.layout_aware_extractor import LayoutAwareExtractor

            extractor = LayoutAwareExtractor()
            result = extractor.extract_with_layout_awareness(file_path)

            main_text = (result or {}).get("text", "")
            tables_text = (result or {}).get("tables_text", "")

            if main_text:
                text_parts.append(main_text)
            if tables_text:
                text_parts.append("\n--- TABLES ---\n" + tables_text)

            combined = "\n\n".join(text_parts).strip()
            if self._has_meaningful_text(combined):
                logger.info(
                    "Extracted PDF using layout-aware method: %s",
                    (result or {}).get("extraction_method", "unknown"),
                )
                return combined
            logger.warning("Layout-aware extraction returned limited text, trying hybrid extractor")
        except Exception as e:
            logger.warning(f"Layout-aware PDF extraction failed: {e}")

        try:
            from app.services.hybrid_extractor import HybridLayoutExtractor

            hybrid = HybridLayoutExtractor()
            result = hybrid.extract_hybrid(file_path)
            hybrid_text = (result or {}).get("full_text", "").strip()
            if self._has_meaningful_text(hybrid_text):
                logger.info("Extracted PDF using hybrid method: %s", (result or {}).get("extraction_method"))
                return hybrid_text
            logger.warning("Hybrid extraction returned limited text, trying brute-force OCR")
        except Exception as e:
            logger.warning(f"Hybrid PDF extraction failed: {e}")

        try:
            brute_text = self._brute_force_ocr_pdf(file_path).strip()
            if self._has_meaningful_text(brute_text):
                logger.info("Extracted PDF using brute-force OCR fallback")
                return brute_text
        except Exception as e:
            logger.warning(f"Brute-force OCR extraction failed: {e}")

        logger.info("Falling back to legacy PDF extraction strategy")
        return self._extract_from_pdf(file_path)

    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type from extension or content"""
        ext = Path(file_path).suffix.lower()

        type_map = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".doc": "doc",
            ".txt": "txt",
            ".png": "png",
            ".jpg": "jpg",
            ".jpeg": "jpeg",
            ".tiff": "tiff",
            ".bmp": "bmp",
        }

        return type_map.get(ext, "unknown")

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using multiple methods"""
        text_parts = []

        if self.is_screenshot_based(file_path):
            logger.info("Detected screenshot-based PDF, prioritizing OCR extraction")
            try:
                return self._extract_from_image(file_path)
            except Exception as e:
                logger.warning(f"OCR-first extraction failed for screenshot-based PDF: {e}")

        # Method 1: pdfplumber (good for structured PDFs)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            if text_parts:
                logger.info("Successfully extracted text with pdfplumber")
                return "\n".join(text_parts)
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")

        # Method 2: PyMuPDF (fallback)
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()

            if text_parts:
                logger.info("Successfully extracted text with PyMuPDF")
                return "\n".join(text_parts)
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}")

        # Method 3: OCR for scanned PDFs
        logger.info("Attempting OCR on PDF")
        return self._extract_from_image(file_path)

    def is_screenshot_based(self, file_path: str) -> bool:
        """Detect whether PDF is mostly image-based with very low native text."""
        try:
            image_count = 0
            text_count = 0

            with fitz.open(file_path) as doc:
                for page in doc:
                    image_count += len(page.get_images(full=True) or [])
                    text_count += len((page.get_text("text") or "").strip())

            return image_count > 3 and text_count < 500
        except Exception:
            return False

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

    def _extract_from_docx_with_images(self, file_path: str) -> str:
        """Extract DOCX text including OCR from embedded images."""
        text_parts: list[str] = []
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                para = (paragraph.text or "").strip()
                if para:
                    text_parts.append(para)

            with zipfile.ZipFile(file_path, "r") as docx_zip:
                for member in docx_zip.namelist():
                    lower = member.lower()
                    if not lower.startswith("word/media/"):
                        continue
                    if not lower.endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif")):
                        continue

                    try:
                        image_bytes = docx_zip.read(member)
                        image_text = self._ocr_image_robust(image_bytes).strip()
                        if image_text:
                            text_parts.append(f"[Image Content] {image_text}")
                    except Exception as ex:
                        logger.debug("Skipping DOCX image %s due to OCR error: %s", member, ex)

            if text_parts:
                return "\n".join(text_parts)
            return self._extract_from_docx(file_path)
        except Exception as e:
            logger.warning(f"DOCX extraction with images failed, falling back: {e}")
            return self._extract_from_docx(file_path)

    def _extract_from_doc(self, file_path: str) -> str:
        """Extract text from legacy DOC files"""
        logger.warning("Legacy .doc extraction not fully implemented")
        try:
            import textract

            text = textract.process(file_path).decode("utf-8")
            return text
        except ImportError:
            raise NotImplementedError(
                "Legacy .doc extraction requires textract. "
                "Install with: pip install textract"
            )
        except Exception as e:
            logger.error(f"DOC extraction failed: {e}")
            raise

    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            encodings = ["latin-1", "cp1252", "iso-8859-1"]
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode file with any encoding")

    def _extract_from_image(self, file_path: str) -> str:
        """Extract text from images using OCR with backend fallback."""
        errors = []

        if self.ocr:
            try:
                result = self._run_ocr(file_path)
                text_parts = self._extract_text_lines(result)
                text = "\n".join(text_parts).strip()
                if text:
                    return text
                errors.append("PaddleOCR returned empty text")
            except Exception as e:
                errors.append(f"PaddleOCR failed: {e}")
                logger.warning(f"PaddleOCR extraction failed; falling back if available: {e}")

        if TESSERACT_AVAILABLE:
            try:
                text = self._extract_with_tesseract(file_path)
                if text:
                    return text
                errors.append("Tesseract returned empty text")
            except Exception as e:
                errors.append(f"Tesseract failed: {e}")
                logger.warning(f"Tesseract extraction failed: {e}")

        raise RuntimeError(
            "Image OCR failed. "
            + " | ".join(errors)
            + " | Install/configure at least one OCR backend (PaddleOCR or pytesseract+tesseract)."
        )

    def _extract_with_tesseract(self, file_path: str) -> str:
        """
        Tesseract fallback OCR for image formats and scanned PDFs.
        Requires both `pytesseract` python package and Tesseract binary in PATH.
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            text_parts = []
            with fitz.open(file_path) as doc:
                for page in doc:
                    pix = page.get_pixmap(dpi=300)
                    img = Image.open(BytesIO(pix.tobytes("png")))
                    text_parts.append(self._tesseract_image_to_text(img))
            return "\n".join([p for p in text_parts if p]).strip()

        with Image.open(file_path) as img:
            return self._tesseract_image_to_text(img)

    def _brute_force_ocr_pdf(self, file_path: str) -> str:
        """Last resort: render all PDF pages and OCR each page."""
        page_texts: list[str] = []
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc):
                extracted = ""
                for zoom in (3.0, 2.5, 2.0):
                    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                    extracted = self._ocr_image_robust(pix.tobytes("png")).strip()
                    if extracted:
                        break
                if not extracted:
                    extracted = (page.get_text("text") or "").strip()
                if extracted:
                    page_texts.append(f"--- Page {page_num + 1} ---\n{extracted}")
        return "\n\n".join(page_texts).strip()

    def _ocr_image_robust(self, img_data: bytes) -> str:
        """OCR bytes with preprocessing + backend fallback."""
        image = Image.open(BytesIO(img_data))
        candidates: list[Image.Image] = [image, image.convert("L")]

        # Optional OpenCV preprocessing variants.
        try:
            import cv2
            import numpy as np

            gray = cv2.cvtColor(np.array(image.convert("RGB")), cv2.COLOR_RGB2GRAY)
            denoised = cv2.fastNlMeansDenoising(gray, h=30)
            _, otsu = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            adaptive = cv2.adaptiveThreshold(
                gray,
                255,
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY,
                11,
                2,
            )
            candidates.append(Image.fromarray(otsu))
            candidates.append(Image.fromarray(adaptive))
        except Exception:
            pass

        best = ""
        if TESSERACT_AVAILABLE:
            for candidate in candidates:
                try:
                    text = pytesseract.image_to_string(candidate, lang="eng").strip()
                    if len(text) > len(best):
                        best = text
                except Exception:
                    continue
            if best:
                return best

        # Paddle fallback using temp image file.
        if self.ocr:
            temp_path = None
            try:
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp:
                    candidates[-1].save(temp, format="PNG")
                    temp_path = temp.name
                result = self._run_ocr(temp_path)
                text_parts = self._extract_text_lines(result)
                return "\n".join(text_parts).strip()
            except Exception:
                return best
            finally:
                if temp_path and os.path.exists(temp_path):
                    try:
                        os.remove(temp_path)
                    except Exception:
                        pass

        return best

    def _tesseract_image_to_text(self, image: Image.Image) -> str:
        """Run Tesseract with light pre-processing for better OCR stability."""
        gray = image.convert("L")
        text = pytesseract.image_to_string(gray, lang="eng")
        return text.strip()

    def _run_ocr(self, file_path: str):
        """
        Run OCR across PaddleOCR API variants.
        Old versions use ocr(..., cls=True); newer builds may reject cls or favor predict(...).
        """
        try:
            return self.ocr.ocr(file_path, cls=True)
        except TypeError:
            pass

        try:
            return self.ocr.ocr(file_path)
        except TypeError:
            pass

        if hasattr(self.ocr, "predict"):
            return self.ocr.predict(file_path)

        raise RuntimeError("Unsupported PaddleOCR API: neither ocr nor predict worked")

    def _extract_text_lines(self, result) -> list[str]:
        """Normalize text output across PaddleOCR response formats."""
        lines: list[str] = []

        def walk(node):
            if node is None:
                return

            if isinstance(node, dict):
                for key in ("rec_texts", "texts", "text", "ocr_text"):
                    value = node.get(key)
                    if isinstance(value, str):
                        if value.strip():
                            lines.append(value.strip())
                    elif isinstance(value, (list, tuple)):
                        for item in value:
                            if isinstance(item, str) and item.strip():
                                lines.append(item.strip())
                return

            if isinstance(node, (list, tuple)):
                if (
                    len(node) >= 2
                    and isinstance(node[1], (list, tuple))
                    and len(node[1]) >= 1
                    and isinstance(node[1][0], str)
                ):
                    text = node[1][0].strip()
                    if text:
                        lines.append(text)
                    return

                for item in node:
                    walk(item)

        walk(result)
        return lines

    def _clean_text(self, text: str) -> str:
        """Light cleaning of extracted text"""
        if not text:
            return ""

        text = text.replace("\x00", "")
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Preserve line boundaries while normalizing noisy horizontal whitespace.
        text = "\n".join(re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n"))
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = "".join(char for char in text if char == "\n" or char == "\t" or ord(char) >= 32)
        return text.strip()

    def _has_meaningful_text(self, text: str, min_chars: int = 100) -> bool:
        cleaned = (text or "").strip()
        return len(cleaned) >= min_chars
